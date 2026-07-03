"""Export analyses and response drafts to .docx.

Lawyers file and review in Word, so both the structured analysis and the response draft must be
one-click exportable. Produces standard, cleanly-formatted documents (Times New Roman body,
centered REMARKS heading for responses) as in-memory bytes.
"""

from __future__ import annotations

import io

from src.models.schemas import OfficeActionAnalysis, ResponseDraft


def _new_doc():
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    return doc


def analysis_to_docx(analysis: OfficeActionAnalysis) -> bytes:
    """Render an OfficeActionAnalysis as a structured .docx and return the bytes."""
    doc = _new_doc()
    doc.add_heading(f"Office Action Analysis — Application {analysis.application_number}", level=1)
    meta = doc.add_paragraph()
    meta.add_run(
        f"{analysis.rejection_type.title()} rejection · Art Unit {analysis.art_unit or '—'} · "
        f"Examiner {analysis.examiner_name or '—'}"
    ).italic = True

    for rej in analysis.rejections:
        doc.add_heading(
            f"Claim {rej.claim_number} — §{rej.rejection_basis.value} "
            f"({'independent' if rej.is_independent else 'dependent'})",
            level=2,
        )
        refs = ", ".join(r.patent_number for r in rej.cited_references) or "—"
        doc.add_paragraph(f"Cited references: {refs}")
        for m in rej.limitation_mappings:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f'"{m.limitation_text}" ').bold = True
            p.add_run(f"→ {m.mapped_to_reference} ({m.reference_passage})")

    if analysis.unverified_claims:
        doc.add_heading("Flags", level=2)
        for flag in analysis.unverified_claims:
            doc.add_paragraph(flag, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def draft_to_docx(draft: ResponseDraft) -> bytes:
    """Render a ResponseDraft as a USPTO-style response (centered REMARKS heading)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _new_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("REMARKS").bold = True

    doc.add_paragraph(
        f"Response for Application {draft.application_number} (strategy: {draft.strategy})."
    )
    for arg in draft.arguments:
        doc.add_heading(f"Claim {arg.claim_number} — §{arg.rejection_basis.value}", level=3)
        doc.add_paragraph(arg.suggested_amendment or arg.argument_text)
        if arg.supporting_sources:
            src = doc.add_paragraph()
            src.add_run("Sources: " + "; ".join(arg.supporting_sources)).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
